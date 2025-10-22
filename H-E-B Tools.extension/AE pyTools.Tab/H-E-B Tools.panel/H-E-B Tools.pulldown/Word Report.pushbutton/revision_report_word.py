###############################################################################
#  revision_report_word.py  (CPython ≥ 3.8 – 3.13.1)                          #
###############################################################################
#!/usr/bin/env python3
"""
Usage (invoked automatically by Script 1):

    python revision_report_word.py <payload.json> <output.docx>

Reads the JSON payload produced by Script 1 and writes a fully-formatted
Word document using python-docx (install once:  `pip install python-docx`).
"""

import sys, json, os, datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------------------------------------------------- CLI args
try:
    json_in, docx_out = sys.argv[1], sys.argv[2]
except ValueError:
    sys.stderr.write("Expected 2 arguments: payload.json output.docx\n")
    sys.exit(1)

with open(json_in, "r") as fp:
    payload = json.load(fp)

proj   = payload["project"]
tables = payload["revisions"]

# -------------------------------------------------------------------- build doc
doc = Document()

# logo
logo = proj.get("logo_path")
if logo and os.path.isfile(logo):
    doc.add_picture(logo, width=Inches(1.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

# headings & metadata
doc.add_heading("Coolsys Energy Design", level=1)
doc.add_heading("Project Revision Summary", level=2)
meta = doc.add_paragraph()
meta.add_run("Project Number: {}\n".format(proj["number"]))
meta.add_run("Client: {}\n".format(proj["client"]))
meta.add_run("Project Name: {}\n".format(proj["name"]))
meta.add_run("Report Date: {}".format(proj["report_date"]))
doc.add_paragraph()  # spacer

# each revision table
for rev in tables:
    doc.add_heading(rev["header"], level=3)
    cols = rev["columns"]
    table = doc.add_table(rows=1, cols=len(cols), style='Table Grid')
    hdr   = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c

    for row in rev["rows"]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    doc.add_paragraph()  # spacer

# save
doc.save(docx_out)
sys.exit(0)
